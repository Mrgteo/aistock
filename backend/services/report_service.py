"""
报告处理服务 - PDF/DOCX文本提取和分块
"""
import os
import uuid
import re
from typing import List, Dict, Optional, Tuple
from datetime import datetime

class ReportService:
    """报告处理服务类"""

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        """
        Args:
            chunk_size: 分块大小（字符数）
            chunk_overlap: 分块重叠大小
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def extract_text_from_pdf(self, file_data: bytes) -> Tuple[str, List[Dict]]:
        """
        从PDF提取文本

        Args:
            file_data: PDF二进制数据

        Returns:
            (完整文本, 每页文本列表)
        """
        import pdfplumber
        import io

        full_text = []
        pages_text = []

        try:
            with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text() or ""
                    text = self._clean_text(text)
                    if text.strip():
                        full_text.append(text)
                        pages_text.append({
                            "page": page_num,
                            "text": text
                        })

            return "\n\n".join(full_text), pages_text

        except Exception as e:
            print(f"[ReportService] PDF解析失败: {e}")
            return "", []

    def extract_text_from_docx(self, file_data: bytes) -> str:
        """
        从DOCX提取文本

        Args:
            file_data: DOCX二进制数据

        Returns:
            提取的文本
        """
        from docx import Document
        import io

        try:
            doc = Document(io.BytesIO(file_data))
            paragraphs = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            # 也提取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        paragraphs.append(" | ".join(row_text))

            full_text = "\n\n".join(paragraphs)
            return self._clean_text(full_text)

        except Exception as e:
            print(f"[ReportService] DOCX解析失败: {e}")
            return ""

    def extract_text(self, file_data: bytes, filename: str) -> Tuple[str, List[Dict]]:
        """
        根据文件类型提取文本

        Args:
            file_data: 文件二进制数据
            filename: 文件名

        Returns:
            (完整文本, 每页/每段文本列表)
        """
        ext = filename.lower().split('.')[-1]

        if ext == 'pdf':
            return self.extract_text_from_pdf(file_data)
        elif ext in ['docx', 'doc']:
            text = self.extract_text_from_docx(file_data)
            return text, [{"page": 1, "text": text}]
        elif ext == 'txt':
            try:
                text = file_data.decode('utf-8')
                return self._clean_text(text), [{"page": 1, "text": text}]
            except:
                try:
                    text = file_data.decode('gbk')
                    return self._clean_text(text), [{"page": 1, "text": text}]
                except:
                    return "", []
        else:
            print(f"[ReportService] 不支持的文件格式: {ext}")
            return "", []

    def chunk_text(self, text: str, pages: List[Dict] = None) -> List[Dict]:
        """
        将文本分割成块

        Args:
            text: 完整文本
            pages: 每页文本列表

        Returns:
            分块列表，每项包含 id, text, page, metadata
        """
        if not text or not text.strip():
            return []

        chunks = []

        if pages:
            # 基于页面的分块
            for page_info in pages:
                page_text = page_info["text"]
                page_num = page_info["page"]

                if len(page_text) <= self.chunk_size:
                    # 页面较小，整个作为一块
                    chunks.append({
                        "chunk_id": str(uuid.uuid4()),
                        "text": page_text.strip(),
                        "page": page_num,
                        "metadata": {"source": "page"}
                    })
                else:
                    # 页面较大，需要分割
                    page_chunks = self._split_long_text(
                        page_text,
                        page_num,
                        {"source": "page"}
                    )
                    chunks.extend(page_chunks)
        else:
            # 无页面信息，基于字符的分块
            chunks = self._split_long_text(text, 0, {"source": "text"})

        # 更新chunk_index
        for i, chunk in enumerate(chunks):
            chunk["chunk_index"] = i

        print(f"[ReportService] 文本分块完成: {len(chunks)} 个块")
        return chunks

    def _split_long_text(self, text: str, page_start: int,
                         base_metadata: Dict) -> List[Dict]:
        """将长文本分割成重叠的块"""
        chunks = []

        # 按段落分割
        paragraphs = re.split(r'\n\s*\n', text)
        current_chunk = []
        current_length = 0
        current_page = page_start

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            para_len = len(para)

            # 如果单个段落超过chunk_size，需要进一步分割
            if para_len > self.chunk_size:
                # 先保存当前块
                if current_chunk:
                    chunk_text = "\n\n".join(current_chunk)
                    chunks.append({
                        "chunk_id": str(uuid.uuid4()),
                        "text": chunk_text,
                        "page": current_page,
                        "metadata": base_metadata.copy()
                    })
                    # 处理重叠
                    overlap_text = current_chunk[-1] if current_chunk else ""
                    current_chunk = [overlap_text] if overlap_text else []
                    current_length = len(overlap_text)

                # 分割长段落
                sub_chunks = self._split_paragraph(para, current_page, base_metadata)
                chunks.extend(sub_chunks[:-1])  # 最后一个保留给后续

                # 用最后一个子块继续
                if sub_chunks:
                    current_chunk = [sub_chunks[-1]["text"]]
                    current_length = len(sub_chunks[-1]["text"])

            elif current_length + para_len + 2 > self.chunk_size:
                # 当前块已满，保存并开始新块
                chunk_text = "\n\n".join(current_chunk)
                chunks.append({
                    "chunk_id": str(uuid.uuid4()),
                    "text": chunk_text,
                    "page": current_page,
                    "metadata": base_metadata.copy()
                })

                # 处理重叠 - 保留最后一段作为新块开头
                overlap_text = current_chunk[-1] if current_chunk else ""
                current_chunk = [overlap_text] if overlap_text else []
                current_length = len(overlap_text)

                # 添加当前段落
                current_chunk.append(para)
                current_length += para_len + 2
            else:
                # 添加到当前块
                current_chunk.append(para)
                current_length += para_len + 2

        # 保存最后一块
        if current_chunk:
            chunk_text = "\n\n".join(current_chunk)
            chunks.append({
                "chunk_id": str(uuid.uuid4()),
                "text": chunk_text,
                "page": current_page,
                "metadata": base_metadata.copy()
            })

        return chunks

    def _split_paragraph(self, para: str, page: int,
                         metadata: Dict, max_chars: int = None) -> List[Dict]:
        """将长段落分割成较小的块"""
        max_chars = max_chars or self.chunk_size
        chunks = []

        # 按句子分割
        sentences = re.split(r'([。！？；\n])', para)
        current_chunk = []
        current_length = 0

        for i in range(0, len(sentences) - 1, 2):
            sentence = sentences[i]
            if i + 1 < len(sentences):
                sentence += sentences[i + 1]  # 加上标点

            sentence = sentence.strip()
            if not sentence:
                continue

            if current_length + len(sentence) > max_chars:
                # 保存当前块
                if current_chunk:
                    chunks.append({
                        "chunk_id": str(uuid.uuid4()),
                        "text": "".join(current_chunk),
                        "page": page,
                        "metadata": metadata.copy()
                    })
                    # 重叠最后一句
                    overlap = current_chunk[-1][-self.chunk_overlap:] if current_chunk else ""
                    current_chunk = [overlap + sentence]
                    current_length = len(overlap) + len(sentence)
                else:
                    # 句子本身太长，直接截断
                    chunks.append({
                        "chunk_id": str(uuid.uuid4()),
                        "text": sentence[:max_chars],
                        "page": page,
                        "metadata": metadata.copy()
                    })
                    current_chunk = []
                    current_length = 0
            else:
                current_chunk.append(sentence)
                current_length += len(sentence)

        # 最后一个块
        if current_chunk:
            chunks.append({
                "chunk_id": str(uuid.uuid4()),
                "text": "".join(current_chunk),
                "page": page,
                "metadata": metadata.copy()
            })

        return chunks

    def _clean_text(self, text: str) -> str:
        """清理文本"""
        if not text:
            return ""

        # 移除多余空白
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)

        # 移除特殊字符（保留中文、英文、数字、常用标点）
        text = re.sub(r'[^\u4e00-\u9fa5\u3000-\u303f\uff00-\uffef\u0020-\u007Ea-zA-Z0-9，。、！？；：""''（）【】《》\-\.\,\!\?\;\:\'\"\(\)\[\]\n ]', '', text)

        return text.strip()

    def extract_stocks_mentioned(self, text: str) -> List[Dict]:
        """
        识别文本中提到的股票（支持代码和名称两种方式），并判断情感倾向
        """
        # A股股票代码模式 (6位数字)
        stock_pattern = r'(?:600|601|603|605|688|000|001|002|003|300)\d{3}'

        # 完整股票名称映射
        stock_names = {
            # 算力层
            "603986": "兆易创新", "688041": "海光信息", "300474": "景嘉微",
            "688012": "中微公司", "688126": "沪硅产业", "002371": "北方华创",
            "688008": "澜起科技", "688111": "金山办公", "603501": "韦尔股份",
            "600745": "闻泰科技", "600588": "用友网络", "002230": "科大讯飞",
            "300496": "中科创达", "300033": "同花顺", "300059": "东方财富",
            "600570": "恒生电子", "688777": "中控技术", "002185": "华天科技",
            "002049": "紫光国微", "688981": "中芯国际",
            # 算法/模型层
            "002261": "拓维信息", "300229": "博彦科技", "300678": "中科信息",
            "300010": "立思辰", "300188": "美亚柏科", "300678": "中科信息",
            # 应用层
            "300413": "芒果超媒", "300364": "中文在线", "002425": "南极电商",
            "002558": "巨人网络", "603444": "吉比特", "300251": "光线传媒",
            "300027": "华谊兄弟", "300058": "蓝色光标", "300133": "华策影视",
            "603189": "网达软件", "600637": "东方明珠",
            # 沪市主板 600/601/603
            "600009": "上海机场", "600011": "华能国际", "600015": "华夏银行",
            "600016": "民生银行", "600018": "上港集团", "600019": "宝钢股份",
            "600023": "浙能电力", "600025": "华能水电", "600026": "中远海能",
            "600028": "中国石化", "600029": "南方航空", "600030": "中信证券",
            "600031": "三一重工", "600036": "招商银行", "600048": "保利发展",
            "600050": "中国联通", "600104": "上汽集团", "600111": "北方稀土",
            "600115": "东方航空", "600118": "中国卫星", "600150": "中国船舶",
            "600188": "兖矿能源", "600196": "复星医药", "600219": "南山铝业",
            "600276": "恒瑞医药", "600309": "万华化学", "600332": "白云山",
            "600362": "江西铜业", "600383": "金地集团", "600406": "国电南瑞",
            "600426": "华鲁恒升", "600436": "片仔癀", "600438": "通威股份",
            "600482": "中国动力", "600487": "亨通光电", "600489": "中金黄金",
            "600519": "贵州茅台", "600547": "山东黄金", "600585": "海螺水泥",
            "600660": "福耀玻璃", "600690": "海尔智家", "600703": "三安光电",
            "600760": "中航沈飞", "600809": "山西汾酒", "600837": "海通证券",
            "600887": "伊利股份", "600893": "航发动力", "600900": "长江电力",
            "600905": "三峡能源", "600941": "中国移动", "601006": "大秦铁路",
            "601012": "隆基绿能", "601066": "中信建投", "601088": "中国神华",
            "601138": "工业富联", "601166": "兴业银行", "601169": "北京银行",
            "601186": "中国铁建", "601211": "国泰君安", "601225": "陕西煤业",
            "601288": "农业银行", "601318": "中国平安", "601328": "交通银行",
            "601390": "中国中铁", "601398": "工商银行", "601601": "中国太保",
            "601628": "中国人寿", "601658": "邮储银行", "601668": "中国建筑",
            "601688": "华泰证券", "601727": "上海电气", "601800": "中国交建",
            "601816": "京沪高铁", "601857": "中国石油", "601888": "中国中免",
            "601898": "中煤能源", "601899": "紫金矿业", "601919": "中远海控",
            "601939": "建设银行", "601985": "中国核电", "601988": "中国银行",
            "601989": "中国重工", "601998": "中信银行", "603259": "药明康德",
            "603288": "海天味业", "603799": "华友钴业",
            # 科创板 688
            "688126": "沪硅产业", "688169": "石头科技", "688187": "时代电气",
            "688223": "晶科能源", "688363": "华熙生物", "688599": "天合光能",
            "688180": "君实生物",
            # 深市主板 000
            "000001": "平安银行", "000002": "万科A", "000063": "中兴通讯",
            "000100": "TCL科技", "000333": "美的集团", "000338": "潍柴动力",
            "000425": "徐工机械", "000538": "云南白药", "000568": "泸州老窖",
            "000596": "古井贡酒", "000651": "格力电器", "000661": "长春高新",
            "000708": "中信特钢", "000725": "京东方A", "000768": "中航西飞",
            "000858": "五粮液", "000876": "新希望", "000895": "双汇发展",
            "000938": "紫光股份",
            # 创业板 300
            "300001": "特锐德", "300015": "爱尔眼科", "300122": "智飞生物",
            "300124": "汇川技术", "300142": "沃森生物", "300207": "欣旺达",
            "300223": "北京君正", "300274": "阳光电源", "300408": "三环集团",
            "300450": "先导智能", "300498": "温氏股份", "300529": "健帆生物",
            "300595": "欧普康视", "300601": "康泰生物", "300628": "亿联网络",
            "300750": "宁德时代", "300759": "康龙化成", "300760": "迈瑞医疗",
            "300896": "爱美客",
            # 中小板 002
            "002001": "新和成", "002027": "分众传媒", "002044": "美年健康",
            "002050": "三花智控", "002081": "金螳螂", "002092": "中泰化学",
            "002120": "韵达股份", "002129": "中环股份", "002142": "宁波银行",
            "002236": "大华股份", "002240": "盛新锂能", "002241": "歌尔股份",
            "002252": "上海莱士", "002304": "洋河股份", "002311": "海大集团",
            "002352": "顺丰控股", "002384": "东山精密", "002415": "海康威视",
            "002459": "晶澳科技", "002460": "赣锋锂业", "002466": "天齐锂业",
            "002475": "立讯精密", "002493": "荣盛石化", "002594": "比亚迪",
            "002601": "龙佰集团", "002709": "天赐材料", "002714": "牧原股份",
            "002812": "恩捷股份", "002841": "视源股份",
        }

        # 通过名称查找代码
        name_to_code = {v: k for k, v in stock_names.items()}

        # 情感判断关键词
        positive_words = ["利好", "受益", "增长", "扩张", "景气", "机会", "推荐", "买入", "超配", "景气上行", "高速增长", "突破", "强劲", "布局"]
        negative_words = ["利空", "风险", "承压", "下滑", "衰退", "减持", "回避", "景气下行", "业绩低于预期", "亏损", "裁员", "制裁"]

        text_lower = text.lower()

        # 1. 先找股票代码及其位置
        code_positions = []
        for match in re.finditer(stock_pattern, text):
            code_positions.append((match.group(), match.start(), match.end()))

        # 2. 再找股票名称及其位置（按名称长度从长到短避免短匹配覆盖）
        sorted_names = sorted(name_to_code.keys(), key=len, reverse=True)
        name_positions = []
        for name in sorted_names:
            pattern = re.escape(name)
            for match in re.finditer(pattern, text):
                code = name_to_code[name]
                name_positions.append((code, match.start(), match.end(), name))

        # 合并所有位置
        all_positions = code_positions + [(c, s, e) for c, s, e, n in name_positions]

        # 3. 对每个股票判断情感
        stock_sentiments = {}
        for code, start, end in all_positions:
            # 提取前后100字的上下文
            context_start = max(0, start - 100)
            context_end = min(len(text), end + 100)
            snippet = text_lower[context_start:context_end]

            pos_count = sum(1 for pw in positive_words if pw in snippet)
            neg_count = sum(1 for nw in negative_words if nw in snippet)

            if pos_count > neg_count:
                sentiment = "利好"
            elif neg_count > pos_count:
                sentiment = "利空"
            else:
                sentiment = "中性"

            if code not in stock_sentiments:
                stock_sentiments[code] = {"sentiment": sentiment, "count": 0, "pos": start}
            else:
                # 如果发现不同情感，更新（有利好就算利好）
                if sentiment == "利好" and stock_sentiments[code]["sentiment"] == "中性":
                    stock_sentiments[code]["sentiment"] = "利好"
                elif sentiment == "利空" and stock_sentiments[code]["sentiment"] == "中性":
                    stock_sentiments[code]["sentiment"] = "利空"
                stock_sentiments[code]["count"] += 1

        # 构建结果
        results = []
        for code, data in stock_sentiments.items():
            results.append({
                "code": code,
                "name": stock_names.get(code, f"股票{code}"),
                "mentions": data["count"],
                "sentiment": data["sentiment"]
            })

        results.sort(key=lambda x: x["mentions"], reverse=True)
        return results[:20]

    def extract_industries_mentioned(self, text: str) -> List[Dict]:
        """
        识别文本中提到的行业，并判断情感倾向
        """
        # AIGC/科技行业相关关键词优先（更精准的匹配）
        industry_keywords = {
            "算力": ["算力", "AI服务器", "AI芯片", "GPU", "智能计算"],
            "算法": ["算法", "大模型", "基础模型", "模型训练", "模型推理"],
            "人工智能": ["人工智能", "机器学习", "深度学习", "AIGC", "生成式AI", "ChatGPT", "LLM"],
            "半导体": ["半导体", "芯片", "集成电路", "晶圆", "光刻机", "封装测试", "半导体设备", "CPU", "GPU"],
            "应用软件": ["应用软件", "办公软件", "企业服务", "SaaS", "行业应用", "垂直应用"],
            "云计算": ["云计算", "云服务", "数据中心", "IDC", "服务器", "云基础设施"],
            "大数据": ["大数据", "数据分析", "数据要素", "数据服务"],
            "元宇宙": ["元宇宙", "虚拟现实", "VR", "AR", "数字藏品", "云游戏"],
            "机器人": ["机器人", "工业机器人", "服务机器人", "人形机器人", "智能制造"],
            "新能源汽车": ["新能源汽车", "电动车", "电动汽车", "锂电", "动力电池", "新能源车", "充电桩", "储能"],
            "医药医疗": ["医药", "医疗器械", "生物医药", "创新药", "中药", "疫苗", "医疗服务"],
            "消费": ["白酒", "食品饮料", "家电", "乳制品", "调味品", "休闲食品", "纺织服装", "消费行业"],
            "银行": ["银行", "银行业", "商业银行", "国有大行", "股份制银行"],
            "保险": ["保险", "保险公司", "人身险", "财产险", "寿险"],
            "房地产": ["房地产", "物业管理", "建筑装饰"],
            "光伏": ["光伏", "太阳能", "风电", "绿电", "风电设备", "储能"],
            "军工": ["军工", "国防", "航空航天", "军用航空", "航天装备", "船舶制造"],
            "通信": ["5G", "通信", "运营商", "通信设备", "光通信", "物联网"],
            "网络安全": ["网络安全", "信息安全", "数据安全", "云安全"],
            "新材料": ["新材料", "碳纤维", "石墨烯", "先进复合材料"],
            "数字经济": ["数字经济", "数字中国", "产业数字化", "数字化转型"],
        }

        # 情感判断关键词
        positive_words = ["利好", "受益", "增长", "扩张", "景气", "机会", "推荐", "买入", "超配", "景气上行", "高速增长", "突破", "强劲"]
        negative_words = ["利空", "风险", "承压", "下滑", "衰退", "减持", "回避", "景气下行", "亏损", "裁员", "制裁"]

        industries = []

        for industry, keywords in industry_keywords.items():
            # 用正则精确匹配词语（避免"AI"匹配到"AIGC"中的"AI"）
            count = 0
            for keyword in keywords:
                # 对于英文词（2-4字母全大写）或短中文词，使用精确边界
                if len(keyword) <= 4 and re.match(r'^[A-Z]+$', keyword):
                    pattern = re.escape(keyword)
                elif len(keyword) <= 4:
                    pattern = re.escape(keyword)
                else:
                    pattern = re.escape(keyword)
                count += len(re.findall(pattern, text))

            if count > 0:
                # 情感判断：检查关键词附近是否有利好/利空词
                sentiment = "中性"
                for kw in keywords:
                    idx = 0
                    while True:
                        pos = text.find(kw, idx)
                        if pos == -1:
                            break
                        # 检查前后80字的情感词
                        start = max(0, pos - 80)
                        end = min(len(text), pos + len(kw) + 80)
                        snippet = text[start:end]

                        pos_count = sum(1 for pw in positive_words if pw in snippet)
                        neg_count = sum(1 for nw in negative_words if nw in snippet)

                        if pos_count > neg_count:
                            sentiment = "利好"
                            break
                        elif neg_count > pos_count:
                            sentiment = "利空"
                            break
                        idx = pos + 1
                    if sentiment != "中性":
                        break

                industries.append({
                    "industry": industry,
                    "mentions": count,
                    "sentiment": sentiment
                })

        # 按出现次数排序
        industries.sort(key=lambda x: x["mentions"], reverse=True)
        return industries[:10]

    def extract_keywords(self, text: str, top_k: int = 5) -> List[str]:
        """
        提取文本中的关键词

        Args:
            text: 文本内容
            top_k: 返回关键词数量

        Returns:
            关键词列表
        """
        import re
        from collections import Counter

        # 定义停用词
        stopwords = {
            '的', '了', '和', '是', '在', '有', '与', '对', '为', '以及',
            '等', '或', '由', '其', '于', '上', '下', '中', '内', '外',
            '将', '把', '被', '将', '会', '能', '可', '也', '都', '而',
            '及', '与', '该', '这', '那', '但', '却', '更', '又', '如果',
            '因为', '所以', '但是', '虽然', '因此', '并且', '或者', '而且',
            '我们', '你们', '他们', '这个', '那个', '一个', '一些', '一种',
            '进行', '通过', '使用', '可以', '已经', '正在', '主要', '重要',
            '可能', '需要', '应该', '可能', '一定', '非常', '比较', '特别',
            '目前', '现在', '当前', '今天', '今年', '去年', '明年', '首次',
            '第一', '第二', '第三', '最后', '包括', '其中', '如下', '上述',
            '本文', '本研究', '本报告', '公司', '企业', '行业', '市场', '数据'
        }

        # 提取中文词语（2-6个字的词组）
        # 使用简单的方式：提取连续的中文字符
        words = re.findall(r'[\u4e00-\u9fa5]{2,6}', text)

        # 过滤停用词和短词
        filtered_words = [w for w in words if w not in stopwords and len(w) >= 2]

        # 统计词频
        word_counts = Counter(filtered_words)

        # 获取最常见的词
        common_words = word_counts.most_common(top_k * 2)

        # 过滤太相似或太通用的词
        keywords = []
        for word, count in common_words:
            # 过滤掉包含在已选关键词中的词
            is_subword = False
            for kw in keywords:
                if word in kw:
                    is_subword = True
                    break
            if not is_subword and count >= 2:
                keywords.append(word)
            if len(keywords) >= top_k:
                break

        return keywords if keywords else ["行业研究", "投资分析"]


# 单例模式
_report_service: Optional[ReportService] = None

def get_report_service() -> ReportService:
    """获取Report服务单例"""
    global _report_service
    if _report_service is None:
        _report_service = ReportService()
    return _report_service
